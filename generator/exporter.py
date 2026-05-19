from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
from datetime import datetime

# Highlight colors
COLORS = {
    "decision": RGBColor(0xFF, 0xEB, 0x3B),   # yellow
    "action":   RGBColor(0xC8, 0xE6, 0xC9),   # green
    "risk":     RGBColor(0xFF, 0xCC, 0xBC),   # red/orange
    "key":      RGBColor(0xBB, 0xDE, 0xFB),   # blue
}

def export_docx(content: str, highlights: list, output_dir: str, filename: str) -> dict:
    """
    Export content to .docx with:
    - Highlighted lines (yellow=decisions, green=actions, red=risks, blue=key)
    - Bold section headers
    - Color legend at top
    - User-editable — standard Word format
    """
    try:
        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1.2)
            section.right_margin  = Inches(1.2)

        # Title
        title = doc.add_heading("MeetMind Output", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Generation timestamp
        ts = doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %b %Y, %I:%M %p')}")
        ts.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ts.runs[0].font.size = Pt(9)
        ts.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_paragraph("")  # spacer

        # Legend
        legend_para = doc.add_paragraph()
        legend_para.add_run("Highlight Legend:  ").bold = True
        _add_colored_run(legend_para, " Decisions ", COLORS["decision"])
        legend_para.add_run("  ")
        _add_colored_run(legend_para, " Action Items ", COLORS["action"])
        legend_para.add_run("  ")
        _add_colored_run(legend_para, " Risks/Issues ", COLORS["risk"])
        legend_para.add_run("  ")
        _add_colored_run(legend_para, " Key Points ", COLORS["key"])

        doc.add_paragraph("")  # spacer

        # Build highlight lookup: text → type
        highlight_map = {}
        for h in highlights:
            highlight_map[h["text"].strip()] = h["type"]

        # Process content line by line
        lines = content.split("\n")
        for line in lines:
            stripped = line.strip()

            if not stripped:
                doc.add_paragraph("")
                continue

            # Section headers (lines starting with **)
            if stripped.startswith("**") and stripped.endswith("**"):
                heading_text = stripped.strip("*").strip()
                doc.add_heading(heading_text, level=2)
                continue

            # Table rows (lines with |)
            if stripped.startswith("|") and stripped.endswith("|"):
                _add_table_row(doc, stripped)
                continue

            # Check if this line should be highlighted
            highlight_type = None
            for h_text, h_type in highlight_map.items():
                if h_text in stripped or stripped in h_text:
                    highlight_type = h_type
                    break

            # Add paragraph
            para = doc.add_paragraph()
            if stripped.startswith("- ") or stripped.startswith("* "):
                para.style = "List Bullet"
                text = stripped[2:]
            else:
                text = stripped

            if highlight_type:
                run = para.add_run(text)
                run.font.highlight_color = None  # Word highlight
                # Use background shading via XML directly
                _shade_paragraph(para, COLORS[highlight_type])
            else:
                para.add_run(text)

        # Save
        out_path = os.path.join(output_dir, f"{filename}.docx")
        doc.save(out_path)
        return {"ok": True, "path": out_path}

    except Exception as e:
        return {"ok": False, "error": f"Export failed: {str(e)}"}


def export_txt(content: str, output_dir: str, filename: str) -> dict:
    """Export plain text version."""
    try:
        # Strip markdown bold markers for clean text
        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', content)
        out_path = os.path.join(output_dir, f"{filename}.txt")
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(clean)
        return {"ok": True, "path": out_path}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def _add_colored_run(paragraph, text: str, color: RGBColor):
    """Add a run with background-colored text (simulated with font color)."""
    run = paragraph.add_run(text)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # Add shading to run's parent paragraph via XML
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rpr = run._r.get_or_add_rPr()
    highlight = OxmlElement('w:highlight')
    # map RGB to word highlight name approximately
    name_map = {
        COLORS["decision"]: "yellow",
        COLORS["action"]:   "green",
        COLORS["risk"]:     "darkRed",
        COLORS["key"]:      "cyan",
    }
    highlight.set(qn('w:val'), name_map.get(color, "yellow"))
    rpr.append(highlight)


def _shade_paragraph(para, color: RGBColor):
    """Add background shading to a paragraph via OOXML."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def _add_table_row(doc, line: str):
    """Render a markdown table row as text (Word table too complex for dynamic rows)."""
    cells = [c.strip() for c in line.strip("|").split("|")]
    # Skip separator rows like |---|---|
    if all(set(c) <= {'-', ' ', ':'} for c in cells):
        return
    para = doc.add_paragraph()
    para.add_run("  •  ".join(cells))
    para.runs[0].font.size = Pt(10)
